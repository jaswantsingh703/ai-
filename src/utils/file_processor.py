import os
import logging
import mimetypes
import magic
import tempfile
import PyPDF2
import docx
import pandas as pd
import csv
import json
import xml.etree.ElementTree as ET
from src.security.ai_security import AISecurity

class FileProcessor:
    """
    Utility for processing different types of files.
    Handles various file formats for analysis by the AI assistant.
    """
    
    def __init__(self):
        """Initialize the file processor with security measures."""
        self.security = AISecurity()
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        
        # Allowed file extensions and MIME types
        self.allowed_extensions = {
            # Documents
            '.txt': 'text',
            '.pdf': 'document',
            '.doc': 'document',
            '.docx': 'document',
            '.rtf': 'document',
            '.odt': 'document',
            
            # Data files
            '.csv': 'data',
            '.json': 'data',
            '.xml': 'data',
            '.xlsx': 'spreadsheet',
            '.xls': 'spreadsheet',
            
            # Images
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.gif': 'image',
            '.bmp': 'image',
            
            # Audio
            '.mp3': 'audio',
            '.wav': 'audio',
            '.flac': 'audio',
            '.m4a': 'audio',
            
            # Video
            '.mp4': 'video',
            '.avi': 'video',
            '.mov': 'video',
            '.mkv': 'video',
            
            # Code
            '.py': 'code',
            '.js': 'code',
            '.html': 'code',
            '.css': 'code',
            '.java': 'code',
            '.cpp': 'code',
            '.c': 'code',
            '.h': 'code',
            '.rb': 'code',
            '.php': 'code'
        }
        
        logging.info("FileProcessor initialized")
    
    def validate_file(self, file_path):
        """
        Validate a file for safety and compatibility.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            tuple: (is_valid, file_info or error_message)
        """
        # Check if file exists
        if not os.path.exists(file_path):
            return False, "File not found"
        
        # Check file size
        if os.path.getsize(file_path) > self.max_file_size:
            return False, f"File too large. Maximum size: {self.max_file_size // (1024 * 1024)}MB"
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Check if extension is allowed
        if ext not in self.allowed_extensions:
            return False, f"File type not supported: {ext}"
        
        # Check MIME type for additional security
        try:
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(file_path)
            
            # Verify MIME type matches extension
            expected_mime_category = self.allowed_extensions[ext]
            
            if not self._is_mime_type_valid(mime_type, expected_mime_category):
                logging.warning(f"MIME type mismatch: {mime_type} for file with extension {ext}")
                return False, "File appears to be disguised with incorrect extension"
            
            # Create file info
            file_info = {
                "path": file_path,
                "name": os.path.basename(file_path),
                "size": os.path.getsize(file_path),
                "mime_type": mime_type,
                "extension": ext,
                "category": self.allowed_extensions[ext]
            }
            
            return True, file_info
            
        except Exception as e:
            logging.error(f"Error validating file {file_path}: {e}")
            return False, f"Error validating file: {str(e)}"
    
    def process_file(self, file_path):
        """
        Process a file and extract its content.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: Processed file information and content
        """
        # Validate the file first
        is_valid, result = self.validate_file(file_path)
        
        if not is_valid:
            return {"error": result}
        
        file_info = result
        
        try:
            # Process based on category
            if file_info["category"] == "text":
                content = self._process_text_file(file_path)
            elif file_info["category"] == "document":
                content = self._process_document(file_path, file_info["extension"])
            elif file_info["category"] == "data":
                content = self._process_data_file(file_path, file_info["extension"])
            elif file_info["category"] == "spreadsheet":
                content = self._process_spreadsheet(file_path)
            elif file_info["category"] == "image":
                content = {"image_path": file_path}
            elif file_info["category"] == "audio":
                content = {"audio_path": file_path}
            elif file_info["category"] == "video":
                content = {"video_path": file_path}
            elif file_info["category"] == "code":
                content = self._process_code_file(file_path)
            else:
                content = {"error": "Unsupported file category"}
            
            # Combine file info with content
            result = {**file_info, "content": content}
            return result
            
        except Exception as e:
            logging.error(f"Error processing file {file_path}: {e}")
            return {
                **file_info,
                "error": f"Error processing file: {str(e)}",
                "content": None
            }
    
    def _process_text_file(self, file_path):
        """Process a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
                
                # Truncate if too large
                if len(content) > 100000:
                    content = content[:100000] + "\n[Content truncated due to size]"
                
                return {
                    "text": content,
                    "length": len(content)
                }
        except UnicodeDecodeError:
            # Try again with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1', errors='replace') as file:
                    content = file.read()
                    return {
                        "text": content,
                        "length": len(content)
                    }
            except Exception as e:
                raise Exception(f"Error reading text file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    def _process_document(self, file_path, extension):
        """Process a document file (PDF, DOCX, etc.)."""
        if extension == '.pdf':
            return self._process_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._process_docx(file_path)
        else:
            # Just treat as text for other document types
            return self._process_text_file(file_path)
    
    def _process_pdf(self, file_path):
        """Process a PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                # Extract text from each page
                for i in range(min(num_pages, 50)):  # Limit to 50 pages
                    page = reader.pages[i]
                    text += page.extract_text() + "\n\n"
                
                # Truncate if too large
                if len(text) > 100000:
                    text = text[:100000] + "\n[Content truncated due to size]"
                
                return {
                    "text": text,
                    "num_pages": num_pages,
                    "processed_pages": min(num_pages, 50)
                }
        except Exception as e:
            raise Exception(f"Error processing PDF file: {str(e)}")
    
    def _process_docx(self, file_path):
        """Process a DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Truncate if too large
            if len(text) > 100000:
                text = text[:100000] + "\n[Content truncated due to size]"
            
            return {
                "text": text,
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error processing DOCX file: {str(e)}")
    
    def _process_data_file(self, file_path, extension):
        """Process a data file (CSV, JSON, XML)."""
        if extension == '.csv':
            return self._process_csv(file_path)
        elif extension == '.json':
            return self._process_json(file_path)
        elif extension == '.xml':
            return self._process_xml(file_path)
        else:
            return {"error": f"Unsupported data file extension: {extension}"}
    
    def _process_csv(self, file_path):
        """Process a CSV file."""
        try:
            # First check encoding and delimiter
            encoding = 'utf-8'
            delimiter = ','
            
            # Try to detect encoding and delimiter
            with open(file_path, 'rb') as f:
                sample = f.read(4096)
                
                # Try to determine encoding
                if sample.startswith(b'\xef\xbb\xbf'):
                    encoding = 'utf-8-sig'
                
                # Check for common delimiters
                text_sample = sample.decode(encoding, errors='replace')
                if text_sample.count(';') > text_sample.count(','):
                    delimiter = ';'
                elif text_sample.count('\t') > text_sample.count(','):
                    delimiter = '\t'
            
            # Read the CSV
            df = pd.read_csv(file_path, encoding=encoding, delimiter=delimiter, error_bad_lines=False)
            
            # Get basic stats
            result = {
                "columns": list(df.columns),
                "rows": len(df),
                "column_types": {col: str(df[col].dtype) for col in df.columns},
                "preview": df.head(5).to_dict(orient='records'),
                "delimiter": delimiter,
                "encoding": encoding
            }
            
            return result
        except Exception as e:
            raise Exception(f"Error processing CSV file: {str(e)}")
    
    def _process_json(self, file_path):
        """Process a JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                data = json.load(file)
                
                # Convert to string representation for simple preview
                json_str = json.dumps(data, indent=2)
                preview = json_str[:1000] + "..." if len(json_str) > 1000 else json_str
                
                return {
                    "structure": self._get_json_structure(data),
                    "preview": preview
                }
        except Exception as e:
            raise Exception(f"Error processing JSON file: {str(e)}")
    
    def _process_xml(self, file_path):
        """Process an XML file."""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            structure = {
                "root_tag": root.tag,
                "attributes": root.attrib,
                "children": self._count_xml_children(root)
            }
            
            return structure
        except Exception as e:
            raise Exception(f"Error processing XML file: {str(e)}")
    
    def _process_spreadsheet(self, file_path):
        """Process a spreadsheet file (XLSX, XLS)."""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            result = {
                "sheets": {}
            }
            
            # Process each sheet
            for sheet_name, df in excel_data.items():
                sheet_data = {
                    "columns": list(df.columns),
                    "rows": len(df),
                    "preview": df.head(5).to_dict(orient='records')
                }
                result["sheets"][sheet_name] = sheet_data
            
            return result
        except Exception as e:
            raise Exception(f"Error processing spreadsheet file: {str(e)}")
    
    def _process_code_file(self, file_path):
        """Process a code file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                code = file.read()
                
                return {
                    "code": code,
                    "length": len(code),
                    "lines": code.count('\n') + 1
                }
        except Exception as e:
            raise Exception(f"Error processing code file: {str(e)}")
    
    def _is_mime_type_valid(self, mime_type, expected_category):
        """
        Check if a MIME type matches the expected category.
        
        Args:
            mime_type (str): MIME type
            expected_category (str): Expected category
            
        Returns:
            bool: True if valid, False otherwise
        """
        if expected_category == 'text':
            return mime_type.startswith('text/')
        elif expected_category == 'document':
            return mime_type in [
                'application/pdf',
                'application/msword',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/rtf',
                'application/vnd.oasis.opendocument.text'
            ]
        elif expected_category == 'data':
            return mime_type in [
                'text/csv',
                'application/json',
                'application/xml',
                'text/plain'
            ]
        elif expected_category == 'spreadsheet':
            return mime_type in [
                'application/vnd.ms-excel',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            ]
        elif expected_category == 'image':
            return mime_type.startswith('image/')
        elif expected_category == 'audio':
            return mime_type.startswith('audio/')
        elif expected_category == 'video':
            return mime_type.startswith('video/')
        elif expected_category == 'code':
            return mime_type in [
                'text/plain',
                'text/x-python',
                'text/javascript',
                'text/html',
                'text/css',
                'text/x-java',
                'text/x-c',
                'application/x-httpd-php'
            ]
        return False
    
    def _get_json_structure(self, data, max_depth=3, current_depth=0):
        """
        Extract the structure of a JSON object.
        
        Args:
            data: JSON data
            max_depth (int): Maximum recursion depth
            current_depth (int): Current recursion depth
            
        Returns:
            dict: JSON structure information
        """
        if current_depth >= max_depth:
            return {"type": type(data).__name__, "truncated": True}
            
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": len(data),
                "properties": {k: self._get_json_structure(v, max_depth, current_depth + 1) 
                              for k, v in list(data.items())[:10]}
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "sample": [self._get_json_structure(v, max_depth, current_depth + 1) 
                          for v in data[:5]]
            }
        else:
            return {"type": type(data).__name__}
    
    def _count_xml_children(self, element):
        """
        Count children in an XML element.
        
        Args:
            element: XML element
            
        Returns:
            int: Number of child elements
        """
        return len(element.findall("./"))